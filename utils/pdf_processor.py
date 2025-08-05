import pdfplumber
import re
import streamlit as st
from typing import List, Optional, Tuple
import tempfile
import os
from docx import Document
from PIL import Image
import pytesseract

class PDFProcessor:
    """Handles PDF text extraction and processing"""
    
    def __init__(self):
        self.question_patterns = [
            r'^\d+\.\s+',  # 1. Question
            r'^\d+\)\s+',  # 1) Question
            r'^Q\d+[\.\)]\s+',  # Q1. or Q1) Question
            r'^Question\s+\d+[\.\)]\s+',  # Question 1. or Question 1)
            r'^\(\d+\)\s+',  # (1) Question
        ]
    
    def extract_text_from_document(self, uploaded_file) -> str:
        """Extract text from uploaded document (PDF, Word, or Image)"""
        try:
            # Reset file pointer
            uploaded_file.seek(0)
            file_extension = uploaded_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                return self._extract_text_from_pdf(uploaded_file)
            elif file_extension in ['docx', 'doc']:
                return self._extract_text_from_word(uploaded_file)
            elif file_extension in ['png', 'jpg', 'jpeg']:
                return self._extract_text_from_image(uploaded_file)
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    def _extract_text_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file with enhanced extraction"""
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_path = tmp_file.name
            
            try:
                full_text = ""
                with pdfplumber.open(tmp_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        # Try standard extraction first
                        text = page.extract_text()
                        
                        # If no text, try with different settings
                        if not text or len(text.strip()) < 10:
                            text = page.extract_text(
                                x_tolerance=2,
                                y_tolerance=2,
                                layout=True
                            )
                        
                        # If still no text, try table extraction
                        if not text or len(text.strip()) < 10:
                            tables = page.extract_tables()
                            if tables:
                                for table in tables:
                                    for row in table:
                                        if row:
                                            text += " ".join([cell for cell in row if cell]) + "\n"
                        
                        if text:
                            full_text += f"--- Page {page_num + 1} ---\n{text}\n\n"
                
                return full_text
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
    
    def _extract_text_from_word(self, uploaded_file) -> str:
        """Extract text from Word document"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_path = tmp_file.name
            
            try:
                doc = Document(tmp_path)
                full_text = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        full_text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text += cell.text + " "
                        full_text += "\n"
                
                return full_text
            finally:
                os.unlink(tmp_path)
                
        except Exception as e:
            raise Exception(f"Failed to extract text from Word document: {str(e)}")
    
    def _extract_text_from_image(self, uploaded_file) -> str:
        """Extract text from image using OCR"""
        try:
            # Load image
            image = Image.open(uploaded_file)
            
            # Use OCR to extract text
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                raise Exception("No text found in image. Please ensure the image contains readable text.")
            
            return text
            
        except Exception as e:
            # Provide helpful error message for OCR issues
            if "tesseract" in str(e).lower():
                raise Exception("OCR service not available. Please convert your image to PDF or Word format.")
            else:
                raise Exception(f"Failed to extract text from image: {str(e)}")
    
    def extract_questions(self, text: str) -> List[str]:
        """Extract questions from text using improved pattern matching"""
        
        if not text or not text.strip():
            return []
        
        # Parse questions from the text
        questions = self._parse_questions(text)
        return questions
    

    

    
    def process_reference_documents(self, reference_files) -> str:
        """Process multiple reference files and extract content"""
        
        all_content = ""
        
        for ref_file in reference_files:
            try:
                # Extract text using the universal method
                file_content = self.extract_text_from_document(ref_file)
                
                if file_content.strip():
                    all_content += file_content + "\n\n"
                
            except Exception as e:
                st.warning(f"Could not process reference file {ref_file.name}: {str(e)}")
                continue
        
        return self._chunk_content(all_content)
    
    def _parse_questions(self, text: str) -> List[str]:
        """Enhanced question parsing with better pattern recognition"""
        
        # Debug: Show what text we're working with
        print(f"DEBUG: Text length: {len(text)}")
        print(f"DEBUG: First 500 chars: {text[:500]}")
        
        # Improved question patterns to catch more formats
        enhanced_patterns = [
            r'^\d+\.\s*(.+)',  # 1. Question (with optional space)
            r'^\d+\)\s*(.+)',  # 1) Question  
            r'^Q\d+[\.\)]\s*(.+)',  # Q1. or Q1) Question
            r'^Question\s+\d+[\.\)]\s*(.+)',  # Question 1. Question
            r'^\(\d+\)\s*(.+)',  # (1) Question
            r'^\d+[\.\s]*([A-Z].+)',  # 1. Capital letter start
            r'^[A-Z]\d+[\.\)]\s*(.+)',  # A1. or B1) Question
            r'^\d+\s*[-–]\s*(.+)',  # 1 - Question or 1 – Question
            r'^[a-z]\)\s*(.+)',  # a) Question
            r'^[A-Z]\)\s*(.+)',  # A) Question
            r'^\d+\s+(.{20,})',  # Number followed by space and substantial text
        ]
        
        questions = []
        lines = text.split('\n')
        current_question = ""
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check each pattern
            question_found = False
            for pattern in enhanced_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    # Save previous question if exists
                    if current_question:
                        questions.append(current_question.strip())
                    
                    # Start new question
                    if match.groups():
                        current_question = match.group(1).strip()
                    else:
                        # Remove numbering manually if no groups
                        current_question = re.sub(r'^\d+[\.\)]\s*|^Q\d+[\.\)]\s*|^Question\s+\d+[\.\)]\s*|^\(\d+\)\s*', '', line).strip()
                    
                    question_found = True
                    break
            
            if not question_found and current_question:
                # Check if this line continues the current question
                if not re.match(r'^\d+[\.\)]\s*|^Q\d+[\.\)]\s*|^Question\s+\d+[\.\)]\s*|^\(\d+\)\s*', line):
                    # This line continues the question
                    current_question += " " + line
        
        # Add the last question
        if current_question:
            questions.append(current_question.strip())
        
        print(f"DEBUG: Raw questions found: {len(questions)}")
        for i, q in enumerate(questions):
            print(f"DEBUG Q{i+1}: {q[:100]}...")
        
        # Enhanced filtering and cleaning
        cleaned_questions = []
        for q in questions:
            # Remove extra whitespace and clean up
            q = re.sub(r'\s+', ' ', q).strip()
            
            # Skip if too short (reduced threshold) or just numbers/symbols
            if len(q) < 10 or re.match(r'^[\d\.\)\(\s\-–]+$', q):
                continue
                
            # Skip common non-question text (more lenient)
            skip_patterns = [
                r'^(page|section|chapter|unit)\s+\d+',
                r'^(note|instruction|direction)s?\s*:',
                r'^(answer|solution|hint)\s*:',
                r'^(total|maximum|minimum)\s+marks?',
                r'^(name|roll|class|date)\s*:'
            ]
            
            should_skip = False
            for skip_pattern in skip_patterns:
                if re.match(skip_pattern, q.lower()):
                    should_skip = True
                    break
            
            if not should_skip:
                cleaned_questions.append(q)
        
        # If no questions found with strict patterns, try looser detection
        if not cleaned_questions:
            print("DEBUG: No questions with strict patterns, trying looser detection...")
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if len(line) > 15:
                    # Look for question indicators
                    question_indicators = ['what', 'how', 'why', 'when', 'where', 'which', 'who', 
                                         'define', 'explain', 'describe', 'calculate', 'find', 
                                         'solve', 'prove', 'derive', 'analyze', 'compare', 'discuss']
                    
                    if any(indicator in line.lower() for indicator in question_indicators):
                        cleaned_questions.append(line)
                    elif line.endswith('?'):
                        cleaned_questions.append(line)
        
        print(f"DEBUG: Final cleaned questions: {len(cleaned_questions)}")
        return cleaned_questions
    
    def _chunk_content(self, content: str, max_chunk_size: int = 2000) -> str:
        """Chunk large content for better processing"""
        
        # Split content into sentences
        sentences = re.split(r'[.!?]+', content)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # If adding this sentence would exceed max size, start new chunk
            if len(current_chunk) + len(sentence) > max_chunk_size:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = sentence
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Return first few chunks to avoid overwhelming the AI
        return "\n\n".join(chunks[:5])
